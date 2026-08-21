"""DOCX loader backed by ``python-docx`` (optional extra: ``semfuse[docx]``)."""

from __future__ import annotations

from pathlib import Path

from semfuse.core.exceptions import DocumentLoadError
from semfuse.core.types import Document
from semfuse.language.detector import detect_language


class DocxLoader:
    """Loads a .docx file as a single document (paragraphs joined)."""

    extensions: tuple[str, ...] = (".docx",)

    def load(self, path: str | Path) -> list[Document]:
        try:
            import docx
        except ImportError as exc:
            raise DocumentLoadError(
                "python-docx is required to load DOCX files. "
                "Install it with `pip install semfuse[docx]`."
            ) from exc
        p = Path(path)
        if not p.exists():
            raise DocumentLoadError(f"File not found: {p}")
        try:
            parsed = docx.Document(str(p))
        except Exception as exc:
            raise DocumentLoadError(f"Failed to parse DOCX {p}: {exc}") from exc
        text = "\n\n".join(par.text for par in parsed.paragraphs if par.text.strip())
        if not text.strip():
            return []
        return [
            Document(
                text=text,
                source=str(p),
                title=p.stem,
                language=detect_language(text),
            )
        ]
